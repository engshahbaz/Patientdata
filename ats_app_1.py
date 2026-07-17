import streamlit as st
import PyPDF2
import docx
from docx.shared import Pt, RGBColor
import os
import json
from io import BytesIO
from google import genai
from google.genai import types

# --- Core Backend Parsers ---

def extract_text_from_docx(file):
    doc = docx.Document(file)
    return "\n".join([p.text for p in doc.paragraphs])

def extract_text_from_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    return "".join([page.extract_text() or "" for page in pdf_reader.pages])

def analyze_resume_with_ai(cv_text, jd_text):
    """Performs an advanced, multi-dimensional audit mirroring premium platform metrics."""
    client = genai.Client(api_key="AIzaSyDLOqU3q2P692velSREiNrWGQjiruAhbtA")
    
    prompt = f"""
    You are an elite executive recruiter and AI-driven ATS Auditor. 
    Analyze the following Candidate CV against the provided Job Description (JD).
    
    Provide your assessment strictly in the following JSON format:
    {{
        "overall_score": <integer between 0 and 100>,
        "years_of_experience": "<string>",
        "verdict": "<string>",
        
        "sub_scores": {{
            "ats_parse_rate": <integer 0-100>,
            "quantifying_impact": <integer 0-100>,
            "repetition": <integer 0-100>,
            "grammar_spelling": <integer 0-100>,
            "hr_red_flags": <integer 0-100>
        }},
        
        "issue_logs": {{
            "quantifying_impact": [<list of specific instances lacking metric focus/numbers>],
            "repetition": [<list of repetitive action verbs or phrases used too frequently>],
            "grammar_spelling": [<list of typos or passive syntax corrections>],
            "hr_red_flags": [<list of red flags like missing location context, formatting flaws, or scope mismatches>]
        }},
        
        "gaps_to_100_percent": [<list of missing technical/business skills>],
        "linkedin_headline": "<string>",
        "linkedin_about": "<string>",
        "linkedin_skills": [<list of strings>]
    }}
    
    Job Description: \"\"\"{jd_text}\"\"\"
    Candidate CV: \"\"\"{cv_text}\"\"\"
    """
    try:
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt,
            config=types.GenerateContentConfig(response_mime_type="application/json", temperature=0.1)
        )
        return json.loads(response.text)
    except Exception as e:
        return {"error": f"API Error: {str(e)}"}

def inject_optimization_to_docx(original_file, gaps):
    doc = docx.Document(original_file)
    if doc.paragraphs:
        heading = doc.paragraphs[0].insert_paragraph_before()
    else:
        heading = doc.add_paragraph()
        
    h_run = heading.add_run("🎯 ATS CORE TARGET METRICS & ARCHITECTURE ALIGNMENT")
    h_run.font.bold = True
    h_run.font.size = Pt(13)
    h_run.font.color.rgb = RGBColor(31, 78, 121)
    
    desc = heading.insert_paragraph_before()
    desc.add_run("The following targeted professional proficiencies have been integrated to align explicitly with the infrastructure requirement matrix:").italic = True
    
    for gap in gaps:
        bullet_p = heading.insert_paragraph_before()
        run_bullet = bullet_p.add_run("• ")
        run_bullet.font.bold = True
        run_text = bullet_p.add_run(gap)
        run_text.font.size = Pt(10.5)
        
    divider = heading.insert_paragraph_before()
    divider.add_run("—" * 60).font.color.rgb = RGBColor(180, 180, 180)
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# --- Premium Streamlit UI Layout ---

st.set_page_config(page_title="RecruiterAI Suite Pro", layout="wide", initial_sidebar_state="expanded")

# Advanced Layout Styling
st.markdown("""
    <style>
    .main { background-color: #f8fafc; }
    .stHeading h1 { color: #0f172a; font-weight: 800; }
    div[data-testid="stSidebar"] { background-color: #0f172a; }
    div[data-testid="stSidebar"] * { color: #f1f5f9 !important; }
    
    .card-panel {
        background-color: white;
        padding: 1.5rem;
        border-radius: 12px;
        box-shadow: 0 1px 3px 0 rgba(0, 0, 0, 0.1);
        margin-bottom: 1.25rem;
        border: 1px solid #e2e8f0;
    }
    .score-container {
        background: linear-gradient(135deg, #1e3a8a 0%, #1e1b4b 100%);
        color: white;
        padding: 2rem;
        border-radius: 16px;
        text-align: center;
        box-shadow: 0 4px 20px 0 rgba(30, 58, 138, 0.15);
    }
    .score-num { font-size: 3.5rem; font-weight: 900; color: #38bdf8; margin-bottom: 0px; }
    .score-lbl { font-size: 0.85rem; text-transform: uppercase; color: #94a3b8; letter-spacing: 0.1em; }
    
    .progress-wrapper { margin-bottom: 1rem; background-color: white; padding: 1rem; border-radius: 8px; border: 1px solid #f1f5f9; }
    .progress-title { font-weight: 600; color: #334155; font-size: 0.95rem; display: flex; justify-content: space-between; }
    
    .stTabs [data-baseweb="tab-list"] { gap: 12px; }
    .stTabs [data-baseweb="tab"] {
        background-color: white;
        border: 1px solid #e2e8f0;
        padding: 10px 24px;
        border-radius: 8px;
        font-weight: 600;
    }
    .stTabs [aria-selected="true"] { background-color: #2563eb !important; color: white !important; }
    </style>
""", unsafe_allow_html=True)

# Navigation Controls
with st.sidebar:
    st.image("https://cdn-icons-png.flaticon.com/512/3201/3201521.png", width=55)
    st.markdown("## Operational Controls")
    st.caption("RecruiterAI Audit Platform v3.0")
    st.write("---")
    st.markdown("### Evaluation Context")
    compliance = st.checkbox("Verify Architecture Frameworks (e.g., TOGAF)", value=True)
    st.info("System configured to run comprehensive profile optimization scripts.")

st.title("💼 RecruiterAI Executive Suite")
st.markdown("##### Multi-Layered ATS Content Diagnostics & Professional Identity Matching")
st.write("")

# Layout: Document Workspace
in_col1, in_col2 = st.columns(2)

with in_col1:
    st.markdown('<div class="card-panel">', unsafe_allow_html=True)
    st.markdown("### 📝 1. Job Parameter Targets")
    jd_input = st.text_area("Paste target job metrics/requirements here:", height=220, placeholder="Requirements, responsibilities, technical frameworks...")
    st.markdown('</div>', unsafe_allow_html=True)

with in_col2:
    st.markdown('<div class="card-panel">', unsafe_allow_html=True)
    st.markdown("### 📂 2. Core Profile Document")
    uploaded_file = st.file_uploader("Upload candidate profile artifact (Word or PDF format)", type=["pdf", "docx"])
    st.markdown('<div style="margin-top:14px;"></div>', unsafe_allow_html=True)
    run_audit = st.button("🚀 Run Comprehensive Diagnostic Optimization", use_container_width=True, type="primary")
    st.markdown('</div>', unsafe_allow_html=True)

# Dashboard Analysis Output
if run_audit:
    if not jd_input or not uploaded_file:
        st.error("🚨 Incomplete profile details. Please make sure both fields are filled out before proceeding.")
    else:
        with st.spinner("🔍 Auditing document syntax, parsing text structure, and analyzing impact criteria..."):
            file_ext = uploaded_file.name.split(".")[-1].lower()
            cv_text = extract_text_from_docx(uploaded_file) if file_ext == "docx" else extract_text_from_pdf(uploaded_file)
            
            analysis = analyze_resume_with_ai(cv_text, jd_input)
            
            if "error" in analysis:
                st.error(analysis["error"])
            else:
                st.toast("Diagnostic Matrix Loaded", icon="📊")
                
                # Split View: Main Score Card on Left, Sub-scores on Right
                dash_col1, dash_col2 = st.columns([1, 2])
                
                with dash_col1:
                    o_score = analysis.get("overall_score", 0)
                    st.markdown(f"""
                        <div class="score-container">
                            <div class="score-num">{o_score}/100</div>
                            <div class="score-lbl">Composite ATS Score</div>
                            <div style='margin-top:10px; font-size:0.9rem; color:#e2e8f0;'>Status Vector: <b>{analysis.get('verdict', 'Review')}</b></div>
                            <div style='font-size:0.85rem; color:#94a3b8;'>Experience Found: {analysis.get('years_of_experience', 'N/A')}</div>
                        </div>
                    """, unsafe_allow_html=True)
                
                with dash_col2:
                    sub_scores = analysis.get("sub_scores", {})
                    
                    # Renders progress loops mapping precisely to premium visual elements
                    metrics = [
                        ("🌐 ATS Structural Parse Rate", sub_scores.get("ats_parse_rate", 90)),
                        ("📊 Quantifying Impact Focus", sub_scores.get("quantifying_impact", 70)),
                        ("🔄 Redundant Term Repetition", sub_scores.get("repetition", 80)),
                        ("✏️ Spelling & Syntax Precision", sub_scores.get("grammar_spelling", 90)),
                        ("⚠️ Safety Checks & HR Red Flags", sub_scores.get("hr_red_flags", 85))
                    ]
                    
                    for label, val in metrics:
                        st.markdown(f"""
                            <div class="progress-wrapper">
                                <div class="progress-title"><span>{label}</span> <span>{val}%</span></div>
                            </div>
                        """, unsafe_allow_html=True)
                        st.progress(val / 100)
                
                st.write("")
                st.write("")
                
                # --- CORE WORKSPACE WORKBOOK TABS ---
                t1, t2, t3, t4 = st.tabs([
                    "🔍 Detailed Content Issue Logs", 
                    "🎯 Keyword Gap Analysis",
                    "🛠️ Document Engineering", 
                    "🌐 Identity Branding Panel"
                ])
                
                with t1:
                    st.markdown("### 📋 Line-Item Issue Audits")
                    logs = analysis.get("issue_logs", {})
                    
                    with st.expander("📊 Quantifying Impact Discrepancies", expanded=True):
                        issues = logs.get("quantifying_impact", [])
                        if issues:
                            for item in issues: st.markdown(f"❌ {item}")
                        else:
                            st.success("Excellent metric utilization. No metric deficiencies observed.")
                            
                    with st.expander("🔄 Document Redundancy & Repetition", expanded=False):
                        issues = logs.get("repetition", [])
                        if issues:
                            for item in issues: st.markdown(f"⚠️ {item}")
                        else:
                            st.success("Strong verbal variation. No excessive repetition detected.")
                            
                    with st.expander("✏️ Syntax & Language Flow Diagnostics", expanded=False):
                        issues = logs.get("grammar_spelling", [])
                        if issues:
                            for item in issues: st.markdown(f"📝 {item}")
                        else:
                            st.success("Perfect grammatical score. Formatting structure is sound.")
                            
                    with st.expander("⚠️ System Risks & Executive HR Red Flags", expanded=False):
                        issues = logs.get("hr_red_flags", [])
                        if issues:
                            for item in issues: st.markdown(f"🚩 {item}")
                        else:
                            st.success("No compliance flags or architectural gaps found.")

                with t2:
                    st.markdown("### 🎯 Blueprint Gaps to Achieve 100% Alignment")
                    gaps = analysis.get("gaps_to_100_percent", [])
                    if gaps:
                        g_cols = st.columns(2)
                        for idx, gap in enumerate(gaps):
                            g_cols[idx % 2].error(f"🔍 **Missing Parameter:** {gap}")
                    else:
                        st.success("Complete strategic alignment reached!")

                with t3:
                    st.markdown("### 📄 Real-Time Structural Document Engine")
                    st.write("Click the button below to download the updated, optimized file containing injected keyword alignments:")
                    if file_ext == "docx" and gaps:
                        optimized_doc_bytes = inject_optimization_to_docx(uploaded_file, gaps)
                        st.download_button(
                            label="📥 Download 100% Aligned Word Document",
                            data=optimized_doc_bytes,
                            file_name="Optimized_Enterprise_CV.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    else:
                        st.info("ℹ️ Document injection downloads operate seamlessly with native Word files (.docx) containing pending keyword adjustments.")

                with t4:
                    st.markdown("### 🌐 Executive Web Presence Matching")
                    st.write("Ensure your public professional presentation mirrors your core enterprise documentation:")
                    
                    st.markdown("##### 🏷️ Algorithmically Indexed Profile Headline")
                    st.code(analysis.get("linkedin_headline", "Corporate Infrastructure Executive"), language="")
                    
                    st.markdown("##### 📝 Summary Narrative Profile ('About' Section)")
                    st.markdown(f"<div style='background-color:white; padding:1.25rem; border-radius:8px; border:1px solid #e2e8f0; color:#334155;'>{analysis.get('linkedin_about', '')}</div>", unsafe_allow_html=True)
                    
                    st.write("")
                    st.markdown("##### 🔑 Key Professional Search Term Anchors")
                    l_skills = analysis.get("linkedin_skills", [])
                    s_cols = st.columns(3)
                    for i, sk in enumerate(l_skills):
                        s_cols[i % 3].markdown(f"✔️ `{sk}`")