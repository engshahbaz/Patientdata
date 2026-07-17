import streamlit as st
import PyPDF2
import docx
from docx.shared import Pt, RGBColor
import os
import json
from io import BytesIO
from google import genai
from google.genai import types

# --- Core Backend Parsers ---

def extract_text_from_docx(file):
    doc = docx.Document(file)
    return "\n".join([p.text for p in doc.paragraphs])

def extract_text_from_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    return "".join([page.extract_text() or "" for page in pdf_reader.pages])

def analyze_resume_with_ai(cv_text, jd_text):
    """Deep Semantic Recruitment Audit using Gemini Flash."""
    client = genai.Client(api_key="AIzaSyDLOqU3q2P692velSREiNrWGQjiruAhbtA")
    
    prompt = f"""
    You are an expert executive technical recruiter and elite ATS screening system. 
    Analyze the following Candidate CV against the provided Job Description (JD).
    
    Provide your assessment strictly in the following JSON format:
    {{
        "match_percentage": <integer between 0 and 100>,
        "years_of_experience": "<string>",
        "key_credentials_detected": [<list of strings>],
        "strengths": [<list of strings>],
        "gaps_to_100_percent": [<list of strings>],
        "verdict": "<string>"
    }}
    
    Job Description: \"\"\"{jd_text}\"\"\"
    Candidate CV: \"\"\"{cv_text}\"\"\"
    """
    try:
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt,
            config=types.GenerateContentConfig(response_mime_type="application/json", temperature=0.1)
        )
        return json.loads(response.text)
    except Exception as e:
        return {"error": f"API Error: {str(e)}"}

def inject_optimization_to_docx(original_file, gaps):
    """
    Bulletproof Modifier: Manually appends text changes without using 
    built-in system styles, avoiding template crashes.
    """
    doc = docx.Document(original_file)
    
    # Safely insert a new heading at the absolute top of the document
    if doc.paragraphs:
        heading = doc.paragraphs[0].insert_paragraph_before()
    else:
        heading = doc.add_paragraph()
        
    h_run = heading.add_run("🎯 ATS CORE TARGET METRICS & ARCHITECTURE ALIGNMENT")
    h_run.font.bold = True
    h_run.font.size = Pt(13)
    h_run.font.color.rgb = RGBColor(31, 78, 121) # Professional dark blue
    
    desc = heading.insert_paragraph_before()
    desc.add_run("The following targeted professional proficiencies have been integrated to align explicitly with the infrastructure requirement matrix:").italic = True
    
    # Safely inject the missing attributes manually using standard paragraph strings
    for gap in gaps:
        bullet_p = heading.insert_paragraph_before()
        # Manually prepend a bullet point character to avoid template style missing bugs
        run_bullet = bullet_p.add_run("• ")
        run_bullet.font.bold = True
        
        run_text = bullet_p.add_run(gap)
        run_text.font.size = Pt(10.5)
        
    # Subtle section divider line
    divider = heading.insert_paragraph_before()
    divider.add_run("—" * 60).font.color.rgb = RGBColor(180, 180, 180)
    
    # Save document to memory stream for instant web browser downloading
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# --- Streamlit UI Interface ---

st.set_page_config(page_title="RecruiterAI Ultra Pro", layout="wide")
st.title("💼 RecruiterAI Ultra Pro: Self-Optimizing ATS Engine")
st.subheader("Automate candidate screening and clean document optimization")

col1, col2 = st.columns(2)

with col1:
    st.header("1. Job Description")
    jd_input = st.text_area("Paste target job parameters:", height=300, placeholder="Requirements go here...")

with col2:
    st.header("2. Candidate CV")
    uploaded_file = st.file_uploader("Upload CV (PDF or Word format)", type=["pdf", "docx"])

if st.button("Execute Screening & Auto-Optimization", type="primary"):
    if not jd_input or not uploaded_file:
        st.error("Please provide both the Job Description and Candidate CV.")
    else:
        with st.spinner("Gemini is conducting analysis and drafting text alignment metrics..."):
            file_ext = uploaded_file.name.split(".")[-1].lower()
            cv_text = extract_text_from_docx(uploaded_file) if file_ext == "docx" else extract_text_from_pdf(uploaded_file)
            
            analysis = analyze_resume_with_ai(cv_text, jd_input)
            
            if "error" in analysis:
                st.error(analysis["error"])
            else:
                st.success("Screening Pipeline Completed!")
                
                # Executive Dashboard Metrics
                m_col1, m_col2 = st.columns(2)
                score = analysis.get("match_percentage", 0)
                m_col1.metric("ATS Match Rating", f"{score}%")
                m_col2.metric("Parsed Experience Vector", analysis.get("years_of_experience", "N/A"))
                
                st.write("---")
                
                # Gaps Dashboard 
                st.subheader("🎯 Blueprint to 100% Match (Identified Gaps)")
                gaps = analysis.get("gaps_to_100_percent", [])
                for gap in gaps:
                    st.write(f"⚠️ **Missing Focus:** {gap}")
                
                # File Generation & Download UI
                if file_ext == "docx" and gaps:
                    st.write("---")
                    st.subheader("🛠️ Immediate Automation Download")
                    
                    # Run the crash-proof file modification 
                    optimized_doc_bytes = inject_optimization_to_docx(uploaded_file, gaps)
                    
                    # Renders the downloadable icon block
                    st.download_button(
                        label="📥 Click Here to Download Your 100% Optimized MS Word CV",
                        data=optimized_doc_bytes,
                        file_name="Optimized_Enterprise_CV.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                elif file_ext != "docx":
                    st.warning("ℹ️ To use the real-time file modifier download icon, please upload your CV as a Word Document (.docx).")