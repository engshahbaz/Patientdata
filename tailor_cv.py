import os
import sys
import re
from docx import Document
from pypdf import PdfReader
from google import genai
from google.genai import types

def extract_raw_text_from_pdf(file_path):
    """Extracts raw text content from the base PDF file to feed into the AI model."""
    if not os.path.exists(file_path):
        print(f"Error: The base PDF resume at '{file_path}' was not found.")
        sys.exit(1)
    try:
        reader = PdfReader(file_path)
        full_text = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                full_text.append(text)
        return "\n".join(full_text)
    except Exception as e:
        print(f"Error reading PDF context source: {e}")
        sys.exit(1)

def load_txt_job_description(file_path):
    """Reads the plain text file containing the job description."""
    if not os.path.exists(file_path):
        print(f"Error: The job description file at '{file_path}' was not found.")
        sys.exit(1)
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def get_tailoring_instructions(resume_text, job_description):
    """Queries Gemini to construct an optimized text block based on the source template text."""
    # SECURITY FIX: Pulled dynamically from environment variables to safeguard credentials
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        print("Error: The 'GEMINI_API_KEY' environment variable is missing. Please set it before executing.")
        sys.exit(1)

    client = genai.Client(api_key=api_key)

    system_instruction = (
        "You are an elite executive resume developer specializing in technical ATS optimization. "
        "Your role is to optimize the text content of the user's resume to match the target job description.\n\n"
        "RULES FOR THE MODEL:\n"
        "1. CONCISENESS: Condense the language to fit nicely within a professional, highly readable 5-page framework.\n"
        "2. ATS MATCHING: Weave essential infrastructure, security, and cloud system keywords from the job description directly into summaries and highlights.\n"
        "3. OUTPUT FORMATTING: Provide the rewritten text grouped into paragraphs that can map clean matching blocks back into the original structure. No conversational chatter."
    )

    user_content = f"""
    === TARGET JOB DESCRIPTION ===
    {job_description}

    === MY BASE RESUME TEXT ===
    {resume_text}

    Please produce the optimized text output keeping structural paragraphs grouped.
    """
    
    print("Sending text to Gemini 2.5 Flash for alignment optimization...")
    try:
        config = types.GenerateContentConfig(system_instruction=system_instruction, temperature=0.3)
        response = client.models.generate_content(model="gemini-2.5-flash", contents=user_content, config=config)
        return response.text
    except Exception as e:
        print(f"Gemini API Processing Error: {e}")
        sys.exit(1)

def clean_and_format_paragraph(paragraph_obj, raw_ai_text):
    """Strips raw markdown syntax out of the AI text, clears runs, and applies Word bolding."""
    clean_text = raw_ai_text.strip()
    
    if paragraph_obj.style.name.startswith('List') or paragraph_obj.text.strip().startswith('•'):
        clean_text = re.sub(r'^[\*\-\d\.\s•]+', '', clean_text)
    else:
        clean_text = re.sub(r'^[\*\-\s]+', '', clean_text)

    paragraph_obj.text = ""

    parts = re.split(r'(\*\*.*?\*\*)', clean_text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            bold_text = part.replace('**', '')
            run = paragraph_obj.add_run(bold_text)
            run.bold = True
        else:
            if part:
                paragraph_obj.add_run(part)

def apply_text_to_template(template_path, output_path, dynamic_text):
    """Loads base_resume.docx layout template and rewrites text content directly inside structural blocks."""
    if not os.path.exists(template_path):
        print(f"Error: Template design layout file '{template_path}' was not found.")
        sys.exit(1)
        
    try:
        doc = Document(template_path)
        ai_paragraphs = [p.strip() for p in dynamic_text.split('\n') if p.strip()]
        
        para_idx = 0
        
        # 1. Update document body structural text paragraphs
        for para in doc.paragraphs:
            if para.text.strip() and para_idx < len(ai_paragraphs):
                clean_and_format_paragraph(para, ai_paragraphs[para_idx])
                para_idx += 1

        # 2. Update structural layout tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip() and para_idx < len(ai_paragraphs):
                            clean_and_format_paragraph(para, ai_paragraphs[para_idx])
                            para_idx += 1

        # Save output checking for lock violations
        try:
            doc.save(output_path)
            print(f"\n[✓] Success! Tailored content mapped directly into your template.")
            print(f"Output File Location: {output_path}")
        except PermissionError:
            print(f"\n[✗] Save Error: Unable to overwrite '{output_path}'. Close the file in Word and retry.")
            sys.exit(1)
            
    except Exception as e:
        print(f"Error preserving style mapping layers: {e}")

def main():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    
    # Configuration targets
    pdf_source_path = os.path.join(base_dir, "base_resume.pdf")
    template_docx_path = os.path.join(base_dir, "base_resume.docx")
    job_desc_path = os.path.join(base_dir, "job_description.txt")
    output_path = os.path.join(base_dir, "resume.docx")

    print("Loading local files and PDF documentation structures...")
    resume_raw_text = extract_raw_text_from_pdf(pdf_source_path)
    job_requirements = load_txt_job_description(job_desc_path)

    # Tailoring loop
    optimized_text = get_tailoring_instructions(resume_raw_text, job_requirements)

    # Apply in-place substitution preserving styles
    apply_text_to_template(template_docx_path, output_path, optimized_text)

if __name__ == "__main__":
    main()