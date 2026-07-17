import os
import sys
from docx import Document
from google import genai
from google.genai import types

def load_file_content(file_path):
    """Safely extracts all text content from paragraphs and tables in a docx file."""
    if not os.path.exists(file_path):
        print(f"Error: The file at '{file_path}' was not found.")
        print("Please make sure your base document is named 'base_resume.docx' and is in this folder.")
        sys.exit(1)
        
    try:
        doc = Document(file_path)
        full_text = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
                    
        return "\n".join(full_text)
    except Exception as e:
        print(f"Error reading docx file: {e}")
        sys.exit(1)

def tailor_cv(resume_text, job_description):
    """Sends content to Gemini to tailor for an ATS 100 score and absolute conciseness."""
    
    # Embedded Google Gemini Auth API Key
    api_key = "AQ.Ab8RN6LJzZFv1UxqC2mlKm8td2hPEvUfQBUcGFm-Svxy7Hi7yQ"
    
    # Initialize the modern official Google Gen AI client
    client = genai.Client(api_key=api_key)

    # ATS 100 & Page-Length Constraint System Instructions
    system_instruction = (
        "You are an expert technical executive recruiter specializing in ATS (Applicant Tracking System) optimization. "
        "Your goal is to rewrite the provided resume to match the job description perfectly, achieving a simulated 100 ATS score.\n\n"
        "CRITICAL RULES:\n"
        "1. CONCISENESS & LENGTH: Strip away repetitive phrasing and fluffy filler words. Ensure the density of the text "
        "will naturally translate to a clean Word document of under 5 pages. Focus heavily on impactful, crisp achievements.\n"
        "2. REAL HUMAN READABLE ENGLISH: Avoid generic corporate buzzword stuffing. Write sentences using strong action verbs "
        "and clear, professional language that a human recruiter can read smoothly.\n"
        "3. ATS 100 MATCHING: Identify all core technical keywords, tools, frameworks, and methodologies in the job description. "
        "Weave them directly into the professional summary, core competencies, and existing experience bullet points.\n"
        "4. LAYOUT COMPLIANCE: Use standard, ATS-scannable section titles (e.g., 'Professional Summary', 'Core Competencies', "
        "'Professional Experience', 'Education'). Do not use charts, icons, or text boxes that scramble ATS scanners.\n"
        "5. FACTUAL HONESTY: Absolutely do NOT invent new jobs, companies, dates, or unearned credentials."
    )

    user_content = f"""
    === TARGET JOB DESCRIPTION ===
    {job_description}

    === MY BASE RESUME TEXT ===
    {resume_text}

    Please rewrite my resume to maximize my ATS score to 100, make it concise, human-readable, and formatted for a clean DOCX output.
    """

    print("Sending request to Gemini 2.5 Flash for ATS 100 optimization...")
    
    try:
        config = types.GenerateContentConfig(
            system_instruction=system_instruction,
            temperature=0.3, # Low temperature ensures factual structural fidelity
        )
        
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=user_content,
            config=config
        )
        
        return response.text
        
    except Exception as e:
        print(f"API Error: {e}")
        sys.exit(1)

def write_text_to_docx(text, output_path):
    """Saves the plain tailored text cleanly into a structured docx file."""
    try:
        new_doc = Document()
        lines = text.split("\n")
        
        for line in lines:
            if line.strip():
                # Detect structural headers and make them look clean
                if line.startswith("#") or line.isupper() or "Experience" in line or "Education" in line:
                    clean_line = line.replace("#", "").strip()
                    p = new_doc.add_paragraph()
                    run = p.add_run(clean_line)
                    run.bold = True
                else:
                    new_doc.add_paragraph(line)
            else:
                new_doc.add_paragraph("")
                
        new_doc.save(output_path)
        print(f"[✓] Success! Your ATS-optimized resume has been generated.")
        print(f"Saved directly to: {output_path}")
    except Exception as e:
        print(f"Error writing to target Word file: {e}")

def main():
    # Dynamic folder execution fallback targeting local directory
    base_dir = os.path.dirname(os.path.abspath(__file__))
    
    resume_path = os.path.join(base_dir, "base_resume.docx")
    job_desc_path = os.path.join(base_dir, "job_description.txt")
    output_path = os.path.join(base_dir, "resume.docx")

    print("Loading local data files...")
    resume_text = load_file_content(resume_path)
    job_content = load_file_content(job_desc_path)

    # Run the optimization loop
    tailored_text = tailor_cv(resume_text, job_content)

    # Commit optimized code to resume.docx
    write_text_to_docx(tailored_text, output_path)

if __name__ == "__main__":
    main()